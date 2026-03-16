import re
from collections import Counter

from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models.structure import (
    AnalyzerConfig,
    BlockFeatures,
    DocumentBaseline,
    LogicalBlock,
    StructureAnalysisResult,
)


class StructureAnalyzer:
    MAIN_HEADING_PREFIX_RE = re.compile(
        r"^(?:第[一二三四五六七八九十百千万0-9]+[章节部分]|[一二三四五六七八九十百千万]+、)"
    )
    SUB_HEADING_PREFIX_RE = re.compile(
        r"^(?:[（(][一二三四五六七八九十0-9]+[）)]|\d+(?:\.\d+){1,3}|\d+[.)、]|\(\d+\))"
    )
    LIST_STYLE_RE = re.compile(r"list", re.IGNORECASE)
    DATE_PREFIX_RE = re.compile(
        r"^(?:\d{4}[./-]\d{1,2}(?:[./-]\d{1,2})?|\d{1,2}月\d{1,2}日|\d{1,2}[-~至]\d{1,2}月)"
    )
    CAPTION_RE = re.compile(
        r"^[\[\u3010\(\uff08]?\s*(?:\u56fe|\u8868|Figure|Table)\s*\d+(?:[.-]\d+)*(?:\s*[\uff1a:]|\s*[\]\u3011\)\uff09])?"
    )
    SENTENCE_PUNCT_RE = re.compile(r"[，。；！？?!]")
    CN_CHAR_RE = re.compile(r"[\u4e00-\u9fff]")

    SEMANTIC_LABELS = [
        "main_heading",
        "sub_heading",
        "inline_subheading",
        "body",
        "list_item",
        "caption",
        "unknown",
    ]

    @classmethod
    def analyze(
        cls,
        document: Document,
        config: AnalyzerConfig | None = None,
        debug: bool = False,
    ) -> StructureAnalysisResult:
        cfg = config or AnalyzerConfig()
        blocks = cls._build_blocks(document, cfg)
        baseline = cls._compute_baseline(blocks)

        for i, block in enumerate(blocks):
            prev_block = blocks[i - 1] if i > 0 else None
            next_block = blocks[i + 1] if i + 1 < len(blocks) else None
            label, scores, reasons, body_match, conf = cls._classify_block(
                block, baseline, cfg, prev_block, next_block
            )
            block.semantic_label = label
            block.scores = scores
            block.reasons = reasons
            block.body_baseline_match_score = body_match
            block.final_confidence = conf

        paragraph_semantic_labels, paragraph_location_types = cls._aggregate_paragraph_labels(blocks)

        if debug:
            cls._debug_print(blocks, baseline, paragraph_semantic_labels)

        return StructureAnalysisResult(
            baseline=baseline,
            blocks=blocks,
            paragraph_semantic_labels=paragraph_semantic_labels,
            paragraph_location_types=paragraph_location_types,
        )

    @classmethod
    def _paragraph_key(
        cls,
        location_type: str,
        paragraph_index: int,
        table_path: str | None,
        row_index: int | None,
        col_index: int | None,
    ) -> str:
        if location_type == "table_cell":
            return f"table:{table_path}:r{row_index}:c{col_index}:p{paragraph_index}"
        return f"paragraph:{paragraph_index}"

    @classmethod
    def _build_blocks(cls, document: Document, cfg: AnalyzerConfig) -> list[LogicalBlock]:
        blocks: list[LogicalBlock] = []
        block_id = 1

        for p_index, paragraph in enumerate(document.paragraphs):
            created = cls._paragraph_to_blocks(
                paragraph=paragraph,
                paragraph_index=p_index,
                source_type="paragraph",
                location_type="paragraph",
                table_path=None,
                row_index=None,
                col_index=None,
                cfg=cfg,
                start_block_id=block_id,
            )
            blocks.extend(created)
            block_id += len(created)

        for t_index, table in enumerate(document.tables):
            block_id = cls._collect_table_blocks(
                table=table,
                table_path=str(t_index),
                cfg=cfg,
                blocks=blocks,
                start_block_id=block_id,
            )

        return blocks

    @classmethod
    def _collect_table_blocks(
        cls,
        table,
        table_path: str,
        cfg: AnalyzerConfig,
        blocks: list[LogicalBlock],
        start_block_id: int,
    ) -> int:
        block_id = start_block_id

        for r_index, row in enumerate(table.rows):
            for c_index, cell in enumerate(row.cells):
                for p_index, paragraph in enumerate(cell.paragraphs):
                    created = cls._paragraph_to_blocks(
                        paragraph=paragraph,
                        paragraph_index=p_index,
                        source_type="table_cell_paragraph",
                        location_type="table_cell",
                        table_path=table_path,
                        row_index=r_index,
                        col_index=c_index,
                        cfg=cfg,
                        start_block_id=block_id,
                    )
                    blocks.extend(created)
                    block_id += len(created)

                for nested_i, nested in enumerate(cell.tables):
                    nested_path = f"{table_path}.{r_index}.{c_index}.{nested_i}"
                    block_id = cls._collect_table_blocks(
                        table=nested,
                        table_path=nested_path,
                        cfg=cfg,
                        blocks=blocks,
                        start_block_id=block_id,
                    )

        return block_id

    @classmethod
    def _paragraph_to_blocks(
        cls,
        paragraph,
        paragraph_index: int,
        source_type: str,
        location_type: str,
        table_path: str | None,
        row_index: int | None,
        col_index: int | None,
        cfg: AnalyzerConfig,
        start_block_id: int,
    ) -> list[LogicalBlock]:
        raw = paragraph.text or ""
        if not raw.strip():
            return []

        para_key = cls._paragraph_key(location_type, paragraph_index, table_path, row_index, col_index)

        lines = [raw]
        if cfg.split_soft_breaks:
            lines = [line for line in re.split(r"[\n\v\r]+", raw) if line.strip()]

        blocks: list[LogicalBlock] = []
        bid = start_block_id

        for line_idx, line in enumerate(lines):
            line_text = line.strip()
            if not line_text:
                continue

            segments = cls._split_inline(line_text, paragraph, cfg)
            for split_idx, seg in enumerate(segments):
                clean = cls._clean_text(seg)
                features = cls._extract_features(
                    raw_text=seg,
                    clean_text=clean,
                    paragraph=paragraph,
                    in_table=(location_type == "table_cell"),
                    has_soft_break=(len(lines) > 1),
                )
                blocks.append(
                    LogicalBlock(
                        block_id=bid,
                        paragraph_key=para_key,
                        raw_text=seg,
                        clean_text=clean,
                        paragraph=paragraph,
                        paragraph_index=paragraph_index,
                        line_index=line_idx,
                        split_index=split_idx,
                        source_type=source_type,
                        location_type=location_type,
                        table_path=table_path,
                        row_index=row_index,
                        col_index=col_index,
                        from_soft_break=(len(lines) > 1),
                        from_inline_split=(len(segments) > 1),
                        features=features,
                    )
                )
                bid += 1

        return blocks

    @classmethod
    def _split_inline(cls, line: str, paragraph, cfg: AnalyzerConfig) -> list[str]:
        s = line.strip()
        if not s or not cfg.split_inline_subheading:
            return [s] if s else []

        if cls._protect_heading_line_from_split(s, paragraph):
            return [s]

        if cfg.keep_list_item_integrity and cls._list_like(s, paragraph):
            return [s]

        colon_pos = min([p for p in (s.find("："), s.find(":")) if p >= 0], default=-1)
        if 1 <= colon_pos <= 20:
            left = s[: colon_pos + 1].strip()
            right = s[colon_pos + 1 :].strip()
            if right and len(right) >= 8 and not cls.DATE_PREFIX_RE.match(left):
                return [left, right]

        m = cls.SUB_HEADING_PREFIX_RE.match(s)
        if m:
            tail = s[m.end() :].strip()
            idx = cls._find_body_marker_index(tail, cfg)
            if idx > 0:
                left = (s[: m.end()] + tail[:idx]).strip()
                right = tail[idx:].strip()
                if left and right:
                    return [left, right]

        return [s]

    @classmethod
    def _protect_heading_line_from_split(cls, text: str, paragraph) -> bool:
        clean = cls._clean_text(text)
        if "：" not in clean and ":" not in clean:
            return False

        font_size = cls._paragraph_font_size_pt(paragraph)
        bold_ratio = cls._bold_ratio(paragraph)
        centered = paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
        short = len(clean) <= 40
        large_font = font_size is not None and font_size >= 15

        return short and centered and bold_ratio >= 0.5 and large_font

    @classmethod
    def _find_body_marker_index(cls, text: str, cfg: AnalyzerConfig) -> int:
        indices = []
        for token in cfg.body_marker_tokens:
            i = text.find(token)
            if i >= 0:
                indices.append(i)
        return min(indices) if indices else -1

    @classmethod
    def _extract_features(
        cls,
        raw_text: str,
        clean_text: str,
        paragraph,
        in_table: bool,
        has_soft_break: bool,
    ) -> BlockFeatures:
        style_name = (paragraph.style.name if paragraph.style else "") or ""

        font_size = cls._paragraph_font_size_pt(paragraph)
        bold_ratio = cls._bold_ratio(paragraph)
        italic_ratio = cls._italic_ratio(paragraph)
        alignment = cls._alignment_name(paragraph.alignment)

        text_len = len(clean_text)
        punct_count = len(cls.SENTENCE_PUNCT_RE.findall(clean_text))
        cn_count = len(cls.CN_CHAR_RE.findall(clean_text))

        return BlockFeatures(
            raw_text=raw_text,
            clean_text=clean_text,
            text_length=text_len,
            is_empty=(text_len == 0),
            style_name=style_name,
            in_table=in_table,
            alignment=alignment,
            bold_ratio=bold_ratio,
            italic_ratio=italic_ratio,
            font_size_pt=font_size,
            has_soft_break=has_soft_break,
            has_main_numbering=bool(cls.MAIN_HEADING_PREFIX_RE.match(clean_text)),
            has_sub_numbering=bool(cls.SUB_HEADING_PREFIX_RE.match(clean_text)),
            list_item_like=cls._list_like(clean_text, paragraph),
            caption_like=bool(cls.CAPTION_RE.match(clean_text)),
            ends_with_colon=clean_text.endswith(("：", ":")),
            ends_with_period=clean_text.endswith(("。", ".", "！", "？", "!", "?")),
            chinese_ratio=(cn_count / text_len if text_len else 0.0),
            punctuation_density=(punct_count / text_len if text_len else 0.0),
            long_sentence_like=(text_len >= 28 or punct_count >= 2),
        )

    @classmethod
    def _compute_baseline(cls, blocks: list[LogicalBlock]) -> DocumentBaseline:
        candidates = []
        for b in blocks:
            f = b.features
            if b.location_type != "paragraph":
                continue
            if f.is_empty or f.caption_like or f.list_item_like:
                continue
            if f.has_main_numbering or f.has_sub_numbering:
                continue
            if f.text_length < 12:
                continue
            if f.alignment == "center" and f.text_length <= 30:
                continue
            if cls._is_heading_style_name(f.style_name):
                continue
            if f.bold_ratio >= 0.45 and f.text_length <= 36:
                continue
            if f.font_size_pt is not None and f.font_size_pt >= 16 and f.text_length <= 40:
                continue
            candidates.append(b)

        if not candidates:
            candidates = [b for b in blocks if b.location_type == "paragraph" and not b.features.is_empty]

        font_sizes = [round(b.features.font_size_pt, 1) for b in candidates if b.features.font_size_pt]
        alignments = [b.features.alignment for b in candidates if b.features.alignment]
        lengths = [b.features.text_length for b in candidates if b.features.text_length > 0]
        bolds = [b.features.bold_ratio for b in candidates]

        return DocumentBaseline(
            font_name=cls._most_common_font_name(candidates),
            font_size_pt=cls._mode(font_sizes),
            alignment=cls._mode(alignments),
            paragraph_length=(sum(lengths) / len(lengths) if lengths else None),
            bold_ratio=(sum(bolds) / len(bolds) if bolds else None),
        )

    @classmethod
    def _classify_block(
        cls,
        block: LogicalBlock,
        baseline: DocumentBaseline,
        cfg: AnalyzerConfig,
        prev_block: LogicalBlock | None,
        next_block: LogicalBlock | None,
    ) -> tuple[str, dict[str, float], list[str], float, float]:
        f = block.features
        scores = {label: 0.0 for label in cls.SEMANTIC_LABELS}
        reasons: list[str] = []

        body_match = cls._body_baseline_match_score(f, baseline)
        scores["body"] += 0.8 + body_match

        if f.caption_like:
            scores["caption"] += 4.0
            reasons.append("caption_pattern:+4 caption")

        if f.has_main_numbering:
            scores["main_heading"] += 3.0
            reasons.append("main_numbering:+3 main_heading")

        if f.has_sub_numbering:
            scores["sub_heading"] += 2.4
            reasons.append("sub_numbering:+2.4 sub_heading")

        if f.ends_with_colon:
            scores["inline_subheading"] += 1.8
            reasons.append("ends_with_colon:+1.8 inline_subheading")

        if f.alignment == "center" and f.text_length <= 40:
            scores["main_heading"] += 1.2
            reasons.append("center_short:+1.2 main_heading")

        if f.bold_ratio >= 0.5 and f.text_length <= 40:
            scores["main_heading"] += 0.8
            scores["sub_heading"] += 0.4

        if f.font_size_pt is not None and f.font_size_pt >= 15 and f.text_length <= 50:
            scores["main_heading"] += 0.9

        if f.alignment == "center" and f.bold_ratio >= 0.5 and f.font_size_pt is not None and f.font_size_pt >= 15 and f.text_length <= 50:
            scores["main_heading"] += 0.8
            scores["body"] -= 1.0

        if f.list_item_like:
            scores["list_item"] += 3.6
            scores["main_heading"] -= 1.2
            scores["sub_heading"] -= 0.8
            reasons.append("list_pattern:+3.6 list_item")

        if prev_block and prev_block.features.list_item_like and f.text_length <= 40:
            scores["list_item"] += 0.6
        if next_block and next_block.features.list_item_like and f.text_length <= 40:
            scores["list_item"] += 0.6

        if f.long_sentence_like:
            scores["body"] += 1.1
            scores["main_heading"] -= 1.0
            scores["sub_heading"] -= 0.8

        if baseline.font_size_pt and f.font_size_pt:
            if f.font_size_pt >= baseline.font_size_pt + 1.5:
                scores["main_heading"] += 1.1
                scores["sub_heading"] += 0.7
                reasons.append("font_larger_than_baseline")
            if abs(f.font_size_pt - baseline.font_size_pt) <= 0.8:
                scores["body"] += 0.5

        if baseline.bold_ratio is not None and f.bold_ratio >= baseline.bold_ratio + 0.35:
            scores["sub_heading"] += 0.6

        if baseline.alignment and f.alignment == "center" and baseline.alignment != "center":
            scores["main_heading"] += 0.9

        if f.text_length <= 24:
            scores["sub_heading"] += 0.4

        if block.location_type == "table_cell":
            scores["main_heading"] -= 0.6
            scores["sub_heading"] -= 0.4
            scores["body"] += 0.3
            # Table cells are often short status/name fragments; prefer body unless
            # there is explicit numbering/list/caption evidence.
            if not f.list_item_like and not f.caption_like and not f.has_main_numbering and not f.has_sub_numbering:
                scores["body"] += 1.1

        ranked = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)
        best_label, best_score = ranked[0]
        second_score = ranked[1][1]
        margin = best_score - second_score

        strong_non_body = max(scores["main_heading"], scores["sub_heading"], scores["list_item"], scores["caption"])
        if best_score < cfg.unknown_score_threshold or margin < cfg.unknown_margin_threshold:
            if f.long_sentence_like and not f.caption_like and strong_non_body < scores["body"] + 0.2:
                best_label = "body"
                reasons.append("body_fallback:long_sentence")
            elif not f.has_main_numbering and not f.has_sub_numbering and not f.list_item_like and not f.caption_like:
                if scores["body"] >= 1.4:
                    best_label = "body"
                    reasons.append("body_fallback:generic")
                else:
                    best_label = "unknown"
            else:
                best_label = "unknown"

        confidence = max(0.0, min(1.0, 0.5 + 0.08 * margin + 0.04 * best_score))
        return best_label, scores, reasons, body_match, confidence

    @classmethod
    def _body_baseline_match_score(cls, f: BlockFeatures, baseline: DocumentBaseline) -> float:
        score = 0.0

        if baseline.font_size_pt and f.font_size_pt:
            diff = abs(f.font_size_pt - baseline.font_size_pt)
            if diff <= 0.5:
                score += 0.8
            elif diff <= 1.0:
                score += 0.4
            elif diff >= 2.0:
                score -= 0.4

        if baseline.alignment and f.alignment == baseline.alignment:
            score += 0.4

        if baseline.paragraph_length and baseline.paragraph_length > 0:
            ratio = f.text_length / baseline.paragraph_length
            if 0.7 <= ratio <= 1.6:
                score += 0.6
            elif ratio < 0.35:
                score -= 0.3

        if baseline.bold_ratio is not None:
            if abs(f.bold_ratio - baseline.bold_ratio) <= 0.2:
                score += 0.3
            elif f.bold_ratio > baseline.bold_ratio + 0.5:
                score -= 0.3

        if f.long_sentence_like:
            score += 0.3

        return score

    @classmethod
    def _aggregate_paragraph_labels(
        cls,
        blocks: list[LogicalBlock],
    ) -> tuple[dict[str, str], dict[str, str]]:
        grouped_sem: dict[str, dict[str, float]] = {}
        grouped_loc: dict[str, dict[str, float]] = {}

        for block in blocks:
            key = block.paragraph_key
            grouped_sem.setdefault(key, {})
            grouped_sem[key][block.semantic_label] = grouped_sem[key].get(block.semantic_label, 0.0) + block.final_confidence
            grouped_loc.setdefault(key, {})
            grouped_loc[key][block.location_type] = grouped_loc[key].get(block.location_type, 0.0) + 1.0

        paragraph_sem: dict[str, str] = {}
        paragraph_loc: dict[str, str] = {}
        for key, votes in grouped_sem.items():
            paragraph_sem[key] = max(votes.items(), key=lambda kv: kv[1])[0]
            paragraph_loc[key] = max(grouped_loc[key].items(), key=lambda kv: kv[1])[0]

        return paragraph_sem, paragraph_loc

    @classmethod
    def _list_like(cls, text: str, paragraph) -> bool:
        s = text.strip()
        style_name = (paragraph.style.name if paragraph and paragraph.style else "") or ""
        if cls.LIST_STYLE_RE.search(style_name):
            return True
        if cls.DATE_PREFIX_RE.match(s):
            return True
        if re.match(r"^[-*\u2022]\s+", s):
            return True
        # Simple numbered list items: 1. xxx / 1) xxx / 1\u3001xxx
        # Multi-level numbering like 1.1 / 1.1.1 is handled as sub-headings.
        if re.match(r"^\d+[\)\u3001]\S", s):
            return True
        if re.match(r"^\d+[\)\u3001]\s+", s):
            return True
        if re.match(r"^\d+\.(?!\d)\S", s):
            return True
        if re.match(r"^\d+\.(?!\d)\s+", s):
            return True
        if re.match(r"^[\uff08(]\d+[\uff09)]", s):
            return True
        return False

    @staticmethod
    def _clean_text(text: str) -> str:
        return re.sub(r"\s+", " ", text or "").strip()

    @staticmethod
    def _mode(items: list):
        if not items:
            return None
        return Counter(items).most_common(1)[0][0]

    @staticmethod
    def _most_common_font_name(blocks: list[LogicalBlock]) -> str | None:
        names: list[str] = []
        for block in blocks:
            for run in block.paragraph.runs:
                if run.font.name:
                    names.append(run.font.name)
        return StructureAnalyzer._mode(names)

    @staticmethod
    def _paragraph_font_size_pt(paragraph) -> float | None:
        sizes: list[float] = []
        for run in paragraph.runs:
            if run.font.size is not None:
                sizes.append(run.font.size.pt)
        if sizes:
            return max(sizes)
        if paragraph.style and paragraph.style.font and paragraph.style.font.size:
            return paragraph.style.font.size.pt
        return None

    @staticmethod
    def _bold_ratio(paragraph) -> float:
        return StructureAnalyzer._font_ratio(paragraph, attr="bold")

    @staticmethod
    def _italic_ratio(paragraph) -> float:
        return StructureAnalyzer._font_ratio(paragraph, attr="italic")

    @staticmethod
    def _font_ratio(paragraph, attr: str) -> float:
        total = 0
        enabled = 0
        for run in paragraph.runs:
            count = len(run.text or "")
            if count == 0:
                continue
            total += count
            value = getattr(run, attr, None)
            if value is None and run.font is not None:
                value = getattr(run.font, attr, None)
            if value:
                enabled += count
        if total == 0:
            return 0.0
        return enabled / total

    @staticmethod
    def _alignment_name(alignment) -> str:
        if alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return "center"
        if alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            return "right"
        return "left"

    @staticmethod
    def _is_heading_style_name(style_name: str) -> bool:
        s = (style_name or "").lower()
        return s.startswith("heading") or s.startswith("标题")

    @classmethod
    def _debug_print(
        cls,
        blocks: list[LogicalBlock],
        baseline: DocumentBaseline,
        paragraph_semantic_labels: dict[str, str],
    ) -> None:
        print("[StructureDebug] ---- baseline ----")
        print(baseline)
        print("[StructureDebug] ---- logical blocks ----")
        for block in blocks:
            print(
                f"[StructureDebug] id={block.block_id} src={block.source} "
                f"semantic={block.semantic_label} location={block.location_type} confidence={block.final_confidence:.3f} "
                f"raw={block.raw_text} clean={block.clean_text}"
            )
            print(f"[StructureDebug] body_baseline_match={block.body_baseline_match_score:.3f}")
            print(f"[StructureDebug] features={block.features.to_dict()}")
            print(f"[StructureDebug] scores={block.scores}")
            print(f"[StructureDebug] reasons={block.reasons}")
        print(f"[StructureDebug] paragraph_semantic_labels={paragraph_semantic_labels}")
        print("[StructureDebug] ---- end ----")
