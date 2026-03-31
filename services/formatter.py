import re
from collections import defaultdict
from copy import deepcopy

from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

from models.format_rule import FormatRule, SectionRule
from models.structure import StructureAnalysisResult


class Formatter:
    ALIGNMENT_MAP = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    # 经验换算：中文“1字符”约等于 0.42 cm，避免缩进过大。
    CHAR_TO_CM = 0.42

    @classmethod
    def paragraph_key(cls, location: str, p_index: int, table_path=None, r_index=None, c_index=None) -> str:
        if location == "table_cell":
            return f"table:{table_path}:r{r_index}:c{c_index}:p{p_index}"
        return f"paragraph:{p_index}"

    @classmethod
    def apply(
        cls,
        document: Document,
        rule: FormatRule,
        analysis: StructureAnalysisResult,
    ) -> dict[str, dict[str, object]]:
        if analysis is None:
            raise ValueError("Formatter.apply requires StructureAnalysisResult")

        rule.normalize()
        applied: dict[str, str] = {}
        split_applied: dict[str, bool] = {}
        bold_cleared: dict[str, bool] = {}

        split_targets = cls._build_split_targets(analysis)

        for p_index, paragraph in enumerate(list(document.paragraphs)):
            key = cls.paragraph_key("paragraph", p_index)
            semantic = analysis.paragraph_semantic_labels.get(key, "unknown")
            location = analysis.paragraph_location_types.get(key, "paragraph")
            section_rule = cls._rule_for(semantic, location, rule)
            cls._apply_to_paragraph_with_optional_split(
                paragraph=paragraph,
                key=key,
                semantic=semantic,
                section_rule=section_rule,
                split_plan=split_targets.get(key),
                location=location,
                rule=rule,
                applied=applied,
                split_applied=split_applied,
                bold_cleared=bold_cleared,
            )

        for t_index, table in enumerate(document.tables):
            cls._apply_table_recursive(
                table=table,
                table_path=str(t_index),
                rule=rule,
                analysis=analysis,
                split_targets=split_targets,
                applied=applied,
                split_applied=split_applied,
                bold_cleared=bold_cleared,
            )

        return {
            "applied_rule": applied,
            "paragraph_split_applied": split_applied,
            "runs_bold_cleared": bold_cleared,
        }

    @classmethod
    def _rule_for(cls, semantic: str, location: str, rule: FormatRule) -> SectionRule | None:
        # Table cells primarily follow table rule, with controlled fallback from body
        # for common expectations: no-bold and paragraph indent in body/list content.
        if location == "table_cell":
            needs_bold_fallback = rule.table.bold is None and rule.body.bold is not None
            needs_indent_fallback = (
                semantic in {"body", "list_item"}
                and rule.table.first_line_indent is None
                and rule.body.first_line_indent is not None
            )

            if not needs_bold_fallback and not needs_indent_fallback:
                return rule.table

            merged = SectionRule(**rule.table.to_dict())
            if needs_bold_fallback:
                merged.bold = rule.body.bold
            if needs_indent_fallback:
                merged.first_line_indent = rule.body.first_line_indent
            return merged

        if semantic in {"main_heading", "sub_heading"}:
            return rule.title
        if semantic == "inline_subheading":
            return rule.inline_subheading
        if semantic == "list_item":
            return rule.list_item
        if semantic == "caption":
            return rule.caption
        if semantic == "body":
            return rule.body
        return None

    @classmethod
    def _apply_table_recursive(
        cls,
        table,
        table_path: str,
        rule: FormatRule,
        analysis: StructureAnalysisResult,
        split_targets: dict[str, list[str]],
        applied: dict[str, str],
        split_applied: dict[str, bool],
        bold_cleared: dict[str, bool],
    ) -> None:
        for r_index, row in enumerate(table.rows):
            for c_index, cell in enumerate(row.cells):
                for p_index, paragraph in enumerate(list(cell.paragraphs)):
                    key = cls.paragraph_key("table_cell", p_index, table_path, r_index, c_index)
                    semantic = analysis.paragraph_semantic_labels.get(key, "unknown")
                    location = analysis.paragraph_location_types.get(key, "table_cell")
                    section_rule = cls._rule_for(semantic, location, rule)
                    cls._apply_to_paragraph_with_optional_split(
                        paragraph=paragraph,
                        key=key,
                        semantic=semantic,
                        section_rule=section_rule,
                        split_plan=split_targets.get(key),
                location=location,
                rule=rule,
                        applied=applied,
                        split_applied=split_applied,
                        bold_cleared=bold_cleared,
                    )
                for nested_i, nested in enumerate(cell.tables):
                    nested_path = f"{table_path}.{r_index}.{c_index}.{nested_i}"
                    cls._apply_table_recursive(
                        nested,
                        nested_path,
                        rule,
                        analysis,
                        split_targets,
                        applied,
                        split_applied,
                        bold_cleared,
                    )

    @classmethod
    def _apply_to_paragraph_with_optional_split(
        cls,
        paragraph,
        key: str,
        semantic: str,
        section_rule: SectionRule | None,
        split_plan: list[dict[str, str]] | None,
        location: str,
        rule: FormatRule,
        applied: dict[str, str],
        split_applied: dict[str, bool],
        bold_cleared: dict[str, bool],
    ) -> None:
        split_applied[key] = False
        bold_cleared[key] = False
        if section_rule is None:
            applied[key] = "skip"
            return

        target_paragraphs = [paragraph]
        if split_plan and len(split_plan) > 1:
            split_lines = [item.get("text", "").strip() for item in split_plan if item.get("text", "").strip()]
            if len(split_lines) > 1:
                target_paragraphs = cls._split_paragraph(paragraph, split_lines)
                split_applied[key] = True

        if split_applied[key] and split_plan:
            line_semantics = [item.get("semantic", semantic) for item in split_plan if item.get("text", "").strip()]
            for para, line_semantic in zip(target_paragraphs, line_semantics):
                line_rule = cls._rule_for(line_semantic, location, rule)
                if line_rule is None:
                    continue
                cleared = cls._apply_paragraph_rule(para, line_rule)
                bold_cleared[key] = bold_cleared[key] or cleared
        else:
            for para in target_paragraphs:
                cleared = cls._apply_paragraph_rule(para, section_rule)
                bold_cleared[key] = bold_cleared[key] or cleared
        applied[key] = semantic

    @classmethod
    def _build_split_targets(cls, analysis: StructureAnalysisResult) -> dict[str, list[dict[str, str]]]:
        grouped: dict[str, list] = defaultdict(list)
        for block in analysis.blocks:
            grouped[block.paragraph_key].append(block)

        targets: dict[str, list[dict[str, str]]] = {}
        for key, blocks in grouped.items():
            soft_blocks = [b for b in blocks if b.from_soft_break]
            if not soft_blocks:
                continue

            line_groups: dict[int, list] = defaultdict(list)
            for b in soft_blocks:
                line_groups[b.line_index].append(b)
            if len(line_groups) <= 1:
                continue

            plan: list[dict[str, str]] = []
            for line_idx in sorted(line_groups.keys()):
                segments = sorted(line_groups[line_idx], key=lambda b: b.split_index)
                line_text = " ".join((seg.raw_text or "").strip() for seg in segments).strip()
                if not line_text:
                    continue

                votes: dict[str, float] = {}
                for seg in segments:
                    votes[seg.semantic_label] = votes.get(seg.semantic_label, 0.0) + seg.final_confidence
                line_semantic = max(votes.items(), key=lambda kv: kv[1])[0] if votes else "unknown"
                plan.append({"text": line_text, "semantic": line_semantic})

            if len(plan) > 1:
                targets[key] = plan
        return targets

    @classmethod
    def _split_paragraph(cls, paragraph, lines: list[str]) -> list:
        clean_lines = [re.sub(r"\s+", " ", (line or "")).strip() for line in lines if line and line.strip()]
        if len(clean_lines) <= 1:
            return [paragraph]

        cls._replace_paragraph_text(paragraph, clean_lines[0])
        paragraphs = [paragraph]
        cursor = paragraph
        for text in clean_lines[1:]:
            new_para = cls._insert_paragraph_after(cursor)
            cls._replace_paragraph_text(new_para, text)
            paragraphs.append(new_para)
            cursor = new_para
        return paragraphs

    @classmethod
    def _insert_paragraph_after(cls, paragraph) -> Paragraph:
        new_p = OxmlElement("w:p")
        ppr = paragraph._p.pPr
        if ppr is not None:
            new_p.append(deepcopy(ppr))
        paragraph._p.addnext(new_p)
        return Paragraph(new_p, paragraph._parent)

    @classmethod
    def _replace_paragraph_text(cls, paragraph, text: str) -> None:
        p = paragraph._p
        ppr_tag = qn("w:pPr")
        for child in list(p):
            if child.tag != ppr_tag:
                p.remove(child)
        paragraph.add_run(text)

    @classmethod
    def _apply_paragraph_rule(cls, paragraph, rule: SectionRule) -> bool:
        if not rule.to_dict():
            return False

        if rule.alignment:
            paragraph.alignment = cls.ALIGNMENT_MAP.get(rule.alignment)

        if rule.line_spacing is not None:
            paragraph.paragraph_format.line_spacing = rule.line_spacing

        if rule.first_line_indent is not None:
            paragraph.paragraph_format.first_line_indent = Cm(rule.first_line_indent * cls.CHAR_TO_CM)

        if rule.space_before is not None:
            paragraph.paragraph_format.space_before = Pt(rule.space_before)

        if rule.space_after is not None:
            paragraph.paragraph_format.space_after = Pt(rule.space_after)

        runs_bold_cleared = False
        for run in paragraph.runs:
            if rule.font_name:
                run.font.name = rule.font_name
            if rule.font_size:
                run.font.size = Pt(rule.font_size)
            if rule.bold is not None:
                run.bold = rule.bold
                run.font.bold = rule.bold
                if rule.bold is False:
                    runs_bold_cleared = True
        return runs_bold_cleared
