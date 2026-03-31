import unittest

from docx import Document

from models.structure import AnalyzerConfig
from services.structure_analyzer import StructureAnalyzer


class StructureAnalyzerTests(unittest.TestCase):
    def test_caption_regex_accepts_bracket_and_colon(self) -> None:
        self.assertTrue(StructureAnalyzer.CAPTION_RE.match("[ 图 1："))
        self.assertTrue(StructureAnalyzer.CAPTION_RE.match("图1："))
        self.assertTrue(StructureAnalyzer.CAPTION_RE.match("表 2.3"))

    def test_list_like_distinguishes_multilevel_numbering(self) -> None:
        doc = Document()

        p_multilevel = doc.add_paragraph("1.1 研究背景")
        p_simple_list = doc.add_paragraph("1. 列表项")

        self.assertFalse(StructureAnalyzer._list_like(p_multilevel.text, p_multilevel))
        self.assertTrue(StructureAnalyzer._list_like(p_simple_list.text, p_simple_list))

    def test_table_top_row_short_cells_are_table_header(self) -> None:
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "事项"
        table.cell(0, 1).text = "状态"
        table.cell(1, 0).text = "成本优化"
        table.cell(1, 1).text = "进行中"

        analysis = StructureAnalyzer.analyze(doc, config=AnalyzerConfig(), debug=False)
        header_blocks = [
            b for b in analysis.blocks if b.location_type == "table_cell" and b.row_index == 0
        ]
        self.assertGreaterEqual(len(header_blocks), 2)
        self.assertTrue(all(b.semantic_label == "table_header" for b in header_blocks))

    def test_bracket_placeholder_is_not_heading(self) -> None:
        doc = Document()
        doc.add_paragraph("[ 销售趋势示意图占位 ]")

        analysis = StructureAnalyzer.analyze(doc, config=AnalyzerConfig(), debug=False)
        target = next((b for b in analysis.blocks if "占位" in b.clean_text), None)
        self.assertIsNotNone(target)
        self.assertNotIn(target.semantic_label, {"main_heading", "sub_heading"})


if __name__ == "__main__":
    unittest.main()
